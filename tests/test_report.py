"""Phase 6 tests: pgx_interpreter.report, per PGx_Project_Plan.md Section 5,
Phase 6 (report assembly, sections 1-10, JSON/TSV/HTML/Markdown/docx).

Network-free by construction: report.py never calls evidence.recommend()
or a gene function itself, so these tests build real PGxResults through the
actual call_tpmt/call_dpyd/call_slco1b1 entry points and real VCF fixtures
(same discipline as every other test module here), attach a real
recommendation via evidence.recommend() pointed at the committed evidence
fixtures for the one case that needs it, then hand the results to
build_report().

The to_docx() tests need the optional `python-docx` dependency
(pyproject.toml's [docx] extra) -- guarded with unittest.SkipTest so this
module runs cleanly (with those specific tests reported as SKIP, not
FAIL) in an environment that doesn't have it installed. This is the exact
mechanism pytest documents for skipping a plain test function without
importing pytest, so it needs no special-casing to also work correctly
under real pytest -- see tests/run_tests.py's module docstring.

Plain `assert` statements only -- must run identically under pytest and
tests/run_tests.py (DEVELOPMENT_WORKFLOW.md item 2).
"""
import io
import json
import unittest
from pathlib import Path

try:
    import docx as python_docx

    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


def _require_docx() -> None:
    if not _DOCX_AVAILABLE:
        raise unittest.SkipTest("python-docx not installed (optional [docx] extra)")

from pgx_interpreter import report
from pgx_interpreter.evidence import recommend
from pgx_interpreter.genes.cyp2c19 import call_cyp2c19
from pgx_interpreter.genes.dpyd import call_dpyd
from pgx_interpreter.genes.slco1b1 import call_slco1b1
from pgx_interpreter.genes.tpmt import call_tpmt
from pgx_interpreter.models import GenomeBuild
from pgx_interpreter.normalize import parse_vcf

FIXTURES_TPMT_DIR = Path(__file__).resolve().parent / "fixtures" / "tpmt"
FIXTURES_DPYD_DIR = Path(__file__).resolve().parent / "fixtures" / "dpyd"
FIXTURES_SLCO1B1_DIR = Path(__file__).resolve().parent / "fixtures" / "slco1b1"
FIXTURES_CYP2C19_DIR = Path(__file__).resolve().parent / "fixtures" / "cyp2c19"
FIXTURES_EVIDENCE_DIR = Path(__file__).resolve().parent / "fixtures" / "evidence"


def _tpmt(fixture_name: str, sample_id: str = "TEST"):
    variants = parse_vcf(FIXTURES_TPMT_DIR / fixture_name, GenomeBuild.GRCH38)
    return call_tpmt(variants, sample_id=sample_id, genome_build=GenomeBuild.GRCH38)


def _dpyd(fixture_name: str, sample_id: str = "TEST"):
    variants = parse_vcf(FIXTURES_DPYD_DIR / fixture_name, GenomeBuild.GRCH38)
    return call_dpyd(variants, sample_id=sample_id, genome_build=GenomeBuild.GRCH38)


def _slco1b1(fixture_name: str, sample_id: str = "TEST"):
    variants = parse_vcf(FIXTURES_SLCO1B1_DIR / fixture_name, GenomeBuild.GRCH38)
    return call_slco1b1(variants, sample_id=sample_id, genome_build=GenomeBuild.GRCH38)


def _cyp2c19(fixture_name: str, sample_id: str = "TEST"):
    variants = parse_vcf(FIXTURES_CYP2C19_DIR / fixture_name, GenomeBuild.GRCH38)
    return call_cyp2c19(variants, sample_id=sample_id, genome_build=GenomeBuild.GRCH38)


# --- build_report(): assembly rules ---


def test_build_report_infers_sample_id_from_single_result():
    result = _tpmt("normal_function.vcf", sample_id="HG002")
    rep = report.build_report((result,))
    assert rep.sample_id == "HG002"
    assert rep.results == (result,)


def test_build_report_infers_sample_id_across_agreeing_multi_gene_results():
    results = (_tpmt("normal_function.vcf", "HG002"), _dpyd("normal_function.vcf", "HG002"))
    rep = report.build_report(results)
    assert rep.sample_id == "HG002"
    assert len(rep.results) == 2


def test_build_report_rejects_mismatched_sample_ids():
    results = (_tpmt("normal_function.vcf", "HG002"), _dpyd("normal_function.vcf", "HG003"))
    try:
        report.build_report(results)
        assert False, "expected ValueError for mismatched sample_ids"
    except ValueError as exc:
        assert "HG002" in str(exc) and "HG003" in str(exc)


def test_build_report_rejects_empty_results():
    try:
        report.build_report(())
        assert False, "expected ValueError for empty results"
    except ValueError:
        pass


def test_build_report_accepts_explicit_generated_at_for_reproducibility():
    result = _tpmt("normal_function.vcf")
    rep = report.build_report((result,), generated_at="2026-08-17T00:00:00+00:00")
    assert rep.generated_at == "2026-08-17T00:00:00+00:00"


# --- to_json(): sections 1-10 ---


def test_json_report_has_metadata_and_disclaimer():
    result = _tpmt("normal_function.vcf", "HG002")
    rep = report.build_report((result,), generated_at="2026-08-17T00:00:00+00:00")
    payload = json.loads(report.to_json(rep))
    assert payload["metadata"]["sample_id"] == "HG002"
    assert payload["metadata"]["generated_at"] == "2026-08-17T00:00:00+00:00"
    assert "research/educational software" in payload["metadata"]["disclaimer"]
    assert "not been independently validated for clinical use" in payload["metadata"]["disclaimer"]


def test_json_report_gene_section_matches_supported_tpmt_result():
    result = _tpmt("het_reduced_function.vcf")
    rep = report.build_report((result,))
    payload = json.loads(report.to_json(rep))
    gene = payload["genes"][0]
    assert gene["gene"] == "TPMT"
    assert gene["genome_build"] == "GRCh38"
    assert gene["allele_diplotype_interpretation"]["diplotype"] == "*1/*3C"
    assert gene["predicted_phenotype"]["phenotype"] == "Intermediate Metabolizer"
    assert gene["predicted_phenotype"]["confidence"] == "supported"


def test_json_report_includes_gene_drug_relationship_when_recommendation_attached():
    result = recommend(_tpmt("normal_function.vcf"), cache_dir=FIXTURES_EVIDENCE_DIR)
    rep = report.build_report((result,))
    payload = json.loads(report.to_json(rep))
    gdr = payload["genes"][0]["gene_drug_relationship"]
    assert gdr["drug"] == "azathioprine"
    assert "normal starting dose" in gdr["recommendation_category"]


def test_json_report_gene_drug_relationship_is_explicit_when_no_recommendation():
    # Ambiguous phenotype -> recommend() never attaches anything (evidence.py's
    # own guardrail) -- the report must say so explicitly, not omit the field.
    result = _tpmt("star3a_unphased_ambiguous.vcf")
    rep = report.build_report((result,))
    payload = json.loads(report.to_json(rep))
    gdr = payload["genes"][0]["gene_drug_relationship"]
    assert gdr["drug"] is None
    assert gdr["recommendation_category"] is None


def test_json_report_guideline_source_version_separates_both_evidence_tiers():
    result = recommend(_tpmt("normal_function.vcf"), cache_dir=FIXTURES_EVIDENCE_DIR)
    rep = report.build_report((result,))
    payload = json.loads(report.to_json(rep))
    gsv = payload["genes"][0]["guideline_source_version"]
    assert gsv["allele_definition_source"] == "PharmVar-equivalent (dbSNP-confirmed)"
    assert gsv["phenotype_evidence_source"] == "CPIC (2018 TPMT/NUDT15 guideline, Table 4)"
    assert gsv["recommendation_evidence_source"] == "CPIC via ClinPGx guidelineAnnotation PA166104933"
    # Tier 1 and Tier 2 are independently versioned, not conflated into one
    # field -- Tier 1's version is the CPIC guideline's publication year,
    # Tier 2's is the actual retrieval date, genuinely different axes.
    assert gsv["phenotype_evidence_version"] == "2018"
    assert gsv["recommendation_evidence_version"] == "2026-08-16"


def test_json_report_surfaces_interpretation_notes_for_ambiguous_case():
    result = _tpmt("star3a_unphased_ambiguous.vcf")
    rep = report.build_report((result,))
    payload = json.loads(report.to_json(rep))
    notes = payload["genes"][0]["interpretation_notes"]
    assert len(notes) == 1
    assert "cannot be distinguished without phasing information" in notes[0]


def test_json_report_interpretation_notes_empty_for_clean_call():
    result = _tpmt("normal_function.vcf")
    rep = report.build_report((result,))
    payload = json.loads(report.to_json(rep))
    assert payload["genes"][0]["interpretation_notes"] == []


def test_json_report_limitations_are_real_gene_specific_text():
    result = _tpmt("normal_function.vcf")
    rep = report.build_report((result,))
    payload = json.loads(report.to_json(rep))
    limitations = payload["genes"][0]["limitations"]
    assert len(limitations) >= 3
    assert any("*2" in text and "*3B" in text for text in limitations)


def test_json_report_technical_provenance_covers_all_three_knowledge_sources():
    result = _tpmt("normal_function.vcf")
    rep = report.build_report((result,))
    payload = json.loads(report.to_json(rep))
    provenance = payload["technical_provenance"]
    assert "MIT" in provenance
    assert "CC BY-SA 4.0" in provenance
    assert "DATA_SOURCES_AND_LICENSING.md" in provenance


def test_json_report_multi_gene_report_lists_every_gene():
    results = (
        _tpmt("normal_function.vcf", "HG002"),
        _dpyd("normal_function.vcf", "HG002"),
        _slco1b1("normal_function.vcf", "HG002"),
        _cyp2c19("normal_function.vcf", "HG002"),
    )
    rep = report.build_report(results)
    payload = json.loads(report.to_json(rep))
    # Phase 8: report.py needed zero code changes to support a fourth gene --
    # build_report()/to_json() are driven entirely by each PGxResult's own
    # `gene` field, not a hardcoded gene list (see report.py's _gene_section).
    assert [g["gene"] for g in payload["genes"]] == ["TPMT", "DPYD", "SLCO1B1", "CYP2C19"]


def test_json_report_gene_section_matches_supported_cyp2c19_result():
    result = _cyp2c19("compound_star2_star17.vcf")
    rep = report.build_report((result,))
    payload = json.loads(report.to_json(rep))
    gene = payload["genes"][0]
    assert gene["gene"] == "CYP2C19"
    assert gene["allele_diplotype_interpretation"]["diplotype"] == "*2/*17"
    assert gene["predicted_phenotype"]["phenotype"] == "Intermediate Metabolizer"
    assert gene["predicted_phenotype"]["confidence"] == "supported"
    assert len(gene["interpretation_notes"]) == 1


def test_json_report_includes_cyp2c19_gene_drug_relationship_when_recommendation_attached():
    result = recommend(_cyp2c19("normal_function.vcf"), cache_dir=FIXTURES_EVIDENCE_DIR)
    rep = report.build_report((result,))
    payload = json.loads(report.to_json(rep))
    gdr = payload["genes"][0]["gene_drug_relationship"]
    assert gdr["drug"] == "clopidogrel"
    assert "standard dose (75 mg/day)" in gdr["recommendation_category"]


# --- to_tsv(): tabular summary only ---


def test_tsv_report_has_one_row_per_gene():
    results = (_tpmt("normal_function.vcf", "HG002"), _dpyd("normal_function.vcf", "HG002"))
    rep = report.build_report(results)
    lines = report.to_tsv(rep).strip("\n").split("\n")
    assert len(lines) == 3  # header + 2 gene rows
    header = lines[0].split("\t")
    assert header == report._TSV_FIELDNAMES
    assert "TPMT" in lines[1]
    assert "DPYD" in lines[2]


def test_tsv_report_omits_free_text_sections():
    # Interpretation notes / limitations / technical provenance don't
    # tabulate cleanly -- deliberately not TSV columns (module docstring).
    result = _tpmt("star3a_unphased_ambiguous.vcf")
    rep = report.build_report((result,))
    tsv = report.to_tsv(rep)
    assert "cannot be distinguished without phasing information" not in tsv
    assert "Only 4 variant alleles are recognized" not in tsv


def test_tsv_report_recommendation_columns_populated_when_attached():
    result = recommend(_tpmt("normal_function.vcf"), cache_dir=FIXTURES_EVIDENCE_DIR)
    rep = report.build_report((result,))
    lines = report.to_tsv(rep).strip("\n").split("\n")
    row = dict(zip(lines[0].split("\t"), lines[1].split("\t")))
    assert row["recommended_drug"] == "azathioprine"
    assert "normal starting dose" in row["recommendation_category"]


# --- to_html(): human-readable, all 10 sections ---


def test_html_report_is_well_formed_and_contains_all_sections():
    result = recommend(_tpmt("normal_function.vcf"), cache_dir=FIXTURES_EVIDENCE_DIR)
    rep = report.build_report((result,), generated_at="2026-08-17T00:00:00+00:00")
    document = report.to_html(rep)
    assert document.startswith("<!DOCTYPE html>")
    assert document.rstrip().endswith("</html>")
    for heading in [
        "3. Observed relevant variants",
        "4. Allele / diplotype interpretation",
        "5. Predicted phenotype",
        "6. Relevant gene-drug relationship",
        "7. Guideline source/version",
        "8. Interpretation notes",
        "9. Limitations",
        "10. Technical provenance",
    ]:
        assert heading in document
    assert "not been independently validated for clinical use" in document
    assert "TPMT" in document
    assert "azathioprine" in document


def test_html_report_escapes_disclaimer_quotes_safely():
    # SOFTWARE_DISCLAIMER contains a literal double-quoted phrase
    # ("What this is not") -- confirm it survives html.escape() without
    # breaking the surrounding markup (i.e. no raw '<' introduced).
    result = _tpmt("normal_function.vcf")
    rep = report.build_report((result,))
    document = report.to_html(rep)
    assert "<script" not in document.lower()
    assert document.count("<html") == 1


def test_html_report_multi_gene_includes_every_gene_section():
    results = (
        _tpmt("normal_function.vcf", "HG002"),
        _slco1b1("normal_function.vcf", "HG002"),
        _cyp2c19("normal_function.vcf", "HG002"),
    )
    rep = report.build_report(results)
    document = report.to_html(rep)
    assert document.count('<section class="gene-section">') == 3
    assert "CYP2C19" in document


# --- to_markdown(): same 10 sections, stdlib-only, no optional dependency ---


def test_markdown_report_contains_all_sections_and_disclaimer():
    result = recommend(_tpmt("normal_function.vcf"), cache_dir=FIXTURES_EVIDENCE_DIR)
    rep = report.build_report((result,), generated_at="2026-08-17T00:00:00+00:00")
    document = report.to_markdown(rep)
    assert document.startswith("# PGx Interpretation Report")
    for heading in [
        "### 3. Observed relevant variants",
        "### 4. Allele / diplotype interpretation",
        "### 5. Predicted phenotype",
        "### 6. Relevant gene-drug relationship",
        "### 7. Guideline source/version",
        "### 8. Interpretation notes",
        "### 9. Limitations",
        "## 10. Technical provenance",
    ]:
        assert heading in document
    assert "not been independently validated for clinical use" in document
    assert "TPMT" in document
    assert "azathioprine" in document


def test_markdown_report_renders_a_real_table():
    result = _tpmt("het_reduced_function.vcf")
    rep = report.build_report((result,))
    document = report.to_markdown(rep)
    assert "| Position | REF>ALT | Zygosity | rsID |" in document
    assert "|---|---|---|---|" in document


def test_markdown_report_surfaces_interpretation_notes():
    result = _tpmt("star3a_unphased_ambiguous.vcf")
    rep = report.build_report((result,))
    document = report.to_markdown(rep)
    assert "cannot be distinguished without phasing information" in document


def test_markdown_report_no_recommendation_is_explicit_not_omitted():
    result = _tpmt("star3a_unphased_ambiguous.vcf")
    rep = report.build_report((result,))
    document = report.to_markdown(rep)
    assert "No drug recommendation attached to this result." in document


def test_markdown_report_multi_gene_separates_sections():
    results = (
        _tpmt("normal_function.vcf", "HG002"),
        _dpyd("normal_function.vcf", "HG002"),
        _cyp2c19("normal_function.vcf", "HG002"),
    )
    rep = report.build_report(results)
    document = report.to_markdown(rep)
    assert document.count("## TPMT") == 1
    assert document.count("## DPYD") == 1
    assert document.count("## CYP2C19") == 1
    assert "---" in document  # gene sections are separated


# --- to_docx(): same 10 sections, needs the optional python-docx dependency ---


def test_docx_report_is_a_valid_zip_container():
    _require_docx()
    result = _tpmt("normal_function.vcf")
    rep = report.build_report((result,))
    data = report.to_docx(rep)
    assert isinstance(data, bytes)
    assert data[:4] == b"PK\x03\x04"  # docx is a zip archive


def test_docx_report_round_trips_gene_content_via_python_docx():
    _require_docx()
    result = recommend(_tpmt("normal_function.vcf"), cache_dir=FIXTURES_EVIDENCE_DIR)
    rep = report.build_report((result,))
    data = report.to_docx(rep)
    document = python_docx.Document(io.BytesIO(data))
    body_text = "\n".join(p.text for p in document.paragraphs)
    assert "TPMT" in body_text
    assert "Normal Metabolizer" in body_text
    assert "azathioprine" in body_text


def test_docx_report_disclaimer_is_in_a_shaded_table_not_a_body_paragraph():
    # The disclaimer lives in a table cell (for the shaded-box treatment),
    # not a plain body paragraph -- confirm it's actually reachable via
    # python-docx's own object model, not just present as raw XML text.
    _require_docx()
    result = _tpmt("normal_function.vcf")
    rep = report.build_report((result,))
    data = report.to_docx(rep)
    document = python_docx.Document(io.BytesIO(data))
    assert len(document.tables) >= 1
    disclaimer_cell_text = document.tables[0].rows[0].cells[0].text
    assert "not been independently validated for clinical use" in disclaimer_cell_text


def test_docx_report_multi_gene_has_one_heading_per_gene():
    _require_docx()
    results = (_tpmt("normal_function.vcf", "HG002"), _slco1b1("normal_function.vcf", "HG002"))
    rep = report.build_report(results)
    data = report.to_docx(rep)
    document = python_docx.Document(io.BytesIO(data))
    heading_texts = [p.text for p in document.paragraphs if p.style.name.startswith("Heading 1")]
    assert any(h.startswith("TPMT") for h in heading_texts)
    assert any(h.startswith("SLCO1B1") for h in heading_texts)


def test_docx_report_interpretation_notes_present_as_bullets():
    _require_docx()
    result = _tpmt("two_no_function_alleles.vcf")
    rep = report.build_report((result,))
    data = report.to_docx(rep)
    document = python_docx.Document(io.BytesIO(data))
    bullet_texts = [p.text for p in document.paragraphs if p.style.name == "List Bullet"]
    assert any("phase inferred from genotype dosage" in t for t in bullet_texts)


def test_to_docx_raises_clear_import_error_when_unavailable(monkeypatch=None):
    # Simulate python-docx being unavailable by forcing the internal
    # `import docx` inside to_docx() to fail, and confirm the resulting
    # error is report.py's own descriptive ImportError (with install
    # instructions), not a bare ModuleNotFoundError leaking an internal
    # import line. Implemented via sys.modules manipulation rather than a
    # pytest fixture (monkeypatch) so this still runs under
    # tests/run_tests.py's plain-function discovery.
    import sys

    result = _tpmt("normal_function.vcf")
    rep = report.build_report((result,))

    real_docx_module = sys.modules.pop("docx", None)
    sys.modules["docx"] = None  # forces `import docx` to raise ImportError
    try:
        try:
            report.to_docx(rep)
            assert False, "expected ImportError when 'docx' is unavailable"
        except ImportError as exc:
            assert "python-docx" in str(exc)
            assert "pip install" in str(exc)
    finally:
        del sys.modules["docx"]
        if real_docx_module is not None:
            sys.modules["docx"] = real_docx_module
